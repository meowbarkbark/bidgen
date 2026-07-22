# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "outputs"
OUTPUT = OUTPUT_DIR / "team-share-cost-validation-discussion.docx"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(102, 112, 133)
BLACK = RGBColor(0, 0, 0)
WHITE = RGBColor(255, 255, 255)

HEADER_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
WARNING_FILL = "FFF7ED"
BORDER = "D9DEE7"
BLUE_FILL = "E8EEF5"


def _set_run_fonts(run, latin: str = "Calibri", east_asia: str = "Malgun Gothic") -> None:
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    run.font.name = latin


def set_run_font(
    run,
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    latin: str = "Calibri",
    east_asia: str = "Malgun Gothic",
) -> None:
    _set_run_fonts(run, latin=latin, east_asia=east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size: float, color: RGBColor, bold: bool = False) -> None:
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_fonts.set(qn("w:eastAsia"), "Malgun Gothic")


def set_paragraph_spacing(paragraph, before: float = 0, after: float = 6, line: float = 1.10) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def set_paragraph_border_bottom(paragraph, color: str = "D7DBE2", size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER, size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent: int = TABLE_INDENT_DXA) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.find(qn("w:tblGrid"))
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(1, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx >= len(row.cells):
                continue
            cell = row.cells[idx]
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def mark_repeating_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def style_table(table, widths: list[int], header: bool = True, body_size: float = 9.4) -> None:
    set_table_geometry(table, widths)
    if header and table.rows:
        mark_repeating_header(table.rows[0])
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            if header and row_idx == 0:
                set_cell_shading(cell, HEADER_FILL)
            for paragraph in cell.paragraphs:
                set_paragraph_spacing(paragraph, after=2, line=1.10)
                if col_idx == len(row.cells) - 1 and len(row.cells) > 2 and row_idx > 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_run_font(run, size=body_size, color=INK, bold=(header and row_idx == 0))


def add_para(doc, text: str, size: float = 11, color: RGBColor = INK, bold: bool = False, after: float = 6, before: float = 0):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=before, after=after, line=1.10)
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold)
    return p


def add_heading(doc, text: str, level: int = 1):
    p = doc.add_paragraph(style=f"Heading {level}")
    before = 16 if level == 1 else 12 if level == 2 else 8
    after = 8 if level == 1 else 6 if level == 2 else 4
    set_paragraph_spacing(p, before=before, after=after, line=1.10)
    r = p.add_run(text)
    set_run_font(r, size=16 if level == 1 else 13 if level == 2 else 12, color=BLUE if level < 3 else DARK_BLUE, bold=True)
    return p


def add_callout(doc, title: str, body: str, fill: str = CALLOUT_FILL) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=3, line=1.12)
    title_run = p.add_run(title)
    set_run_font(title_run, size=11, color=DARK_BLUE, bold=True)
    p.add_run("\n")
    body_run = p.add_run(body)
    set_run_font(body_run, size=10.5, color=INK)
    style_table(table, [CONTENT_WIDTH_DXA], header=False)
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, after=4)


def add_table(doc, headers: list[str], rows: list[tuple[str, ...]], widths: list[int], body_size: float = 9.4) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].paragraphs[0].add_run(header)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].paragraphs[0].add_run(value)
    style_table(table, widths, header=True, body_size=body_size)
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, after=4)


def add_label_table(doc, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    for row_idx, (label, value) in enumerate(rows):
        left, right = table.rows[row_idx].cells
        left.paragraphs[0].add_run(label)
        right.paragraphs[0].add_run(value)
        set_cell_shading(left, HEADER_FILL)
    style_table(table, [1800, 7560], header=False, body_size=9.8)
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, after=6)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    set_style_font(normal, size=11, color=INK, bold=False)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = doc.styles[style_name]
        set_style_font(style, size=size, color=color, bold=True)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("원가계산서 자동검증 논의 공유")
    set_run_font(header_run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("Team sharing brief")
    set_run_font(footer_run, size=9, color=MUTED)


def add_masthead(doc: Document) -> None:
    title = doc.add_paragraph()
    set_paragraph_spacing(title, before=0, after=4)
    title_run = title.add_run("원가계산서 자동검증 논의 공유 메모")
    set_run_font(title_run, size=23, color=BLACK, bold=True)

    subtitle = doc.add_paragraph()
    set_paragraph_spacing(subtitle, after=12)
    subtitle_run = subtitle.add_run("Excel 정규화, 기준자료 검증, AI API 역할, Phase 2 공고문 기능 정리")
    set_run_font(subtitle_run, size=12.5, color=MUTED)

    add_label_table(
        doc,
        [
            ("문서 목적", "팀 공유를 위한 대화 요약 및 제품 방향 정리"),
            ("작성 기준", "사용자와 Codex의 2026-07-22 논의"),
            ("현재 결론", "Phase 1은 원가계산서 검증 데이터 생성, Phase 2는 공고문 자동생성"),
            ("관련 산출물", "PRD.md, 청주 검증 HTML 리포트, 표준 원가계산서 정규화 리포트"),
        ],
    )

    rule = doc.add_paragraph()
    set_paragraph_spacing(rule, before=2, after=10)
    set_paragraph_border_bottom(rule)


def build_doc() -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_masthead(doc)

    add_callout(
        doc,
        "한 줄 결론",
        "우리는 공고문 생성기를 바로 만드는 단계가 아니라, 다양한 원가계산서 Excel을 신뢰 가능한 표준 검증 데이터(validatedProject)로 만드는 Phase 1을 먼저 완성하는 방향으로 정리했다.",
    )

    add_heading(doc, "1. 논의 배경", 1)
    add_para(
        doc,
        "처음에는 원가계산서 Excel을 파싱해 고정 JSON 구조로 넣을 수 있는지 확인했다. 이후 샘플 파일들을 보면서 단순 JSON 미리보기만으로는 검수자가 이해하기 어렵고, Excel처럼 표 형태로 확인 가능한 HTML 검수 리포트가 필요하다는 방향으로 확장했다.",
    )
    add_para(
        doc,
        "중간에 기존 6조 기획안의 BidGen 흐름과 현재 작업 범위가 섞여 보였지만, 결론적으로 공고문 기능은 삭제하지 않고 Phase 2로 분리했다. 현재 MVP는 원가계산서 정규화와 검증을 끝까지 믿을 수 있게 만드는 것이 핵심이다.",
    )

    add_heading(doc, "2. 핵심 의사결정", 1)
    add_table(
        doc,
        ["주제", "결정 내용"],
        [
            ("MVP 초점", "Phase 1은 원가계산서 자동검증과 근거 리포트에 집중한다."),
            ("공고문 기능", "이번 MVP에서 제외하되, validatedProject를 입력으로 받는 Phase 2 확장 기능으로 유지한다."),
            ("정규화 산출물", "원본 양식 차이를 흡수해 normalizedCostStatement와 validatedProject JSON으로 정리한다."),
            ("계산 책임", "금액, 요율, 차액 판정은 규칙 엔진이 담당하고 AI가 최종 숫자 판정을 독점하지 않는다."),
            ("AI API 역할", "항목 매핑 후보, 기준자료 적합성 판단, 검증불가 사유 설명, PDF 근거 해석을 보조한다."),
        ],
        [1900, 7460],
    )

    add_heading(doc, "3. 주요 논의 흐름", 1)
    add_table(
        doc,
        ["논의 항목", "정리 내용"],
        [
            ("Excel 정규화", "원가계산서 시트만 우선 표준 컬럼으로 정규화하되, 산출기초가 다른 시트나 수식에 숨어 있으면 추적해 근거값으로 가져온다."),
            ("다른 시트 처리", "모든 시트를 같은 양식으로 정규화할 필요는 없다. 검증에 필요한 기준자료, 산출근거, 세부내역 시트만 목적별 adapter로 추출한다."),
            ("검수용 리포트", "파일 입력 후 로딩이 끝나면 최상단에 검증결과 요약을 보여주고, 아래에는 표준화된 원가계산서와 근거 상세를 보여준다."),
            ("경비 섹션 기준", "원본에서 노무비 소계 다음부터 경비 소계까지가 경비다. 경비 소계 아래 항목은 경비로 밀어 넣지 않는다."),
            ("가변 경비 항목", "경비 항목명은 매번 달라질 수 있으므로 고정 enum이 아니라 sourceLabel, canonicalKey, category, confidence를 가진 유연한 line-item 구조로 둔다."),
            ("기준자료 입력", "제비율표와 노임단가표는 사용자가 Excel로 입력한다. 표준품셈 PDF는 MVP 필수 입력이 아니라 향후 검증 강화용 선택 자료로 둔다."),
            ("공고문 위치", "공고문 생성은 Phase 2에서 validatedProject를 입력받아 자동 초안 생성으로 연결한다."),
        ],
        [2100, 7260],
        body_size=9.1,
    )

    add_heading(doc, "4. 검증에서 확인하는 것", 1)
    add_callout(
        doc,
        "검증의 기준",
        "사용자가 넣는 제비율표와 노임단가표 Excel은 외부 기준값이다. 원가계산서 Excel 내부의 다른 시트는 산출기초, 세부내역, 수식 근거를 추적하는 내부 근거 자료로 본다.",
        fill=BLUE_FILL,
    )
    add_table(
        doc,
        ["검증 대상", "비교 기준", "결과 예시"],
        [
            ("노임단가", "사용자가 업로드한 노임단가표 Excel", "직종별 단가 일치, 불일치, 기준 누락"),
            ("제비율", "사용자가 업로드한 제비율표 Excel", "요율 정상, 조건 확인 필요, 기준 불일치"),
            ("산출기초", "원가계산서의 수식, 참조 셀, 다른 시트 세부내역", "근거 추적 성공, 수식 내 숫자 추출, 근거 없음"),
            ("합계/소계", "정규화된 항목 금액과 원본 소계", "합계 일치, 반올림 차이, 계산 오류"),
            ("항목 매핑", "정규화 규칙과 AI 후보 판단", "표준 항목 매핑, 확인 필요, 검증불가"),
        ],
        [1700, 4000, 3660],
        body_size=9.1,
    )

    add_heading(doc, "5. 정규화가 쓸모 있는 이유", 1)
    add_para(
        doc,
        "정규화는 단순히 보기 좋게 줄 세우는 작업이 아니다. 다양한 원본 양식을 같은 계산 모델로 옮겨야 노임단가, 제비율, 소계, 총계, 산출기초를 같은 방식으로 검증할 수 있다. 즉 정규화가 있어야 검증 로직과 리포트 UI가 파일마다 새로 흔들리지 않는다.",
    )
    add_para(
        doc,
        "다만 정규화는 원본 의미를 보존해야 한다. 이번 논의에서 비파괴검사 항목이 경비로 잘못 들어간 사례처럼, 섹션 경계가 틀리면 계산 자체가 왜곡된다. 그래서 정규화 단계에는 원본 행 번호, 원본 섹션, 추정 confidence, 매핑 근거를 함께 남겨야 한다.",
    )

    add_heading(doc, "6. 청주 샘플 검증에서 확인한 사실", 1)
    add_table(
        doc,
        ["영역", "테스트 결과 요약"],
        [
            ("정규화", "청주 원가계산서 26개 행 매핑 및 섹션 경계 규칙 검증 통과"),
            ("노임단가", "도장공, 보통인부, 특별인부 단가가 업로드한 기준표보다 낮게 입력되어 오류로 감지"),
            ("제비율 정상", "산재보험료, 고용보험료, 퇴직공제부금비, 산업안전보건관리비, 일반관리비, 이윤은 정상 판정"),
            ("확인 필요", "건강보험료, 연금보험료, 노인장기요양보험료는 기준표 요율과 다르지만 1개월 미만 미적용 조건 가능성이 있어 확인 필요"),
            ("검증불가", "간접노무비, 기타경비, 환경보전비 등은 업로드 기준표와 원본 조건이 맞지 않아 자동 확정이 어려움"),
        ],
        [1800, 7560],
        body_size=9.1,
    )

    add_heading(doc, "7. PRD 반영 사항", 1)
    add_table(
        doc,
        ["반영 위치", "변경 내용"],
        [
            ("제품 비전", "BidGen의 최종 방향은 유지하되, 현재 MVP는 원가계산서 검증 특화로 정의"),
            ("범위", "공고문 자동생성은 이번 MVP 제외 및 Phase 2 확장 기능으로 명시"),
            ("데이터 구조", "normalizedCostStatement와 validatedProject를 핵심 산출물로 추가"),
            ("AI API", "AI는 계산 엔진이 아니라 해석/매핑/설명 보조자로 제한"),
            ("검증 흐름", "기준자료 Excel 업로드, 내부 시트 근거 추적, 리포트 기반 확인 흐름 추가"),
            ("향후 로드맵", "Phase 2에서 validatedProject 기반 공고문 초안 생성으로 확장"),
        ],
        [1900, 7460],
    )

    add_heading(doc, "8. 사용자 흐름 요약", 1)
    add_table(
        doc,
        ["단계", "사용자 행동", "시스템 처리", "화면/산출물"],
        [
            ("1", "원가계산서 Excel 업로드", "원가계산서 시트 탐지 및 원본 구조 분석", "업로드 상태"),
            ("2", "노임단가표, 제비율표 Excel 업로드", "외부 기준값 파싱 및 적용 조건 추출", "기준자료 인식 결과"),
            ("3", "검증 실행", "정규화, 수식 추적, 다른 시트 근거 참조, 기준값 비교", "로딩 및 처리 상태"),
            ("4", "검증결과 확인", "오류, 확인 필요, 검증불가, 정상 항목 분류", "상단 요약 + 표 형태 검수 리포트"),
            ("5", "항목별 근거 확인", "원본 셀, 참조 시트, 기준표, 계산식 표시", "상세 근거 패널"),
            ("6", "결과 저장", "validatedProject JSON 생성", "Phase 2 공고문 생성 입력값"),
        ],
        [850, 2350, 3150, 3010],
        body_size=8.7,
    )

    add_heading(doc, "9. 팀 액션 아이템", 1)
    add_table(
        doc,
        ["담당 영역", "액션", "우선순위"],
        [
            ("기획/업무", "Phase 1과 Phase 2의 사용자 스토리를 발표자료와 PRD에 동일하게 정리", "높음"),
            ("데이터/검증", "제비율표의 공종, 금액, 기간 조건 선택 규칙과 검증불가 기준 정리", "높음"),
            ("프론트엔드", "업로드, 검증결과 요약, 표준화 표, 상세 근거 패널 중심 화면 구현", "높음"),
            ("AI/API", "AI 보조 판단 인터페이스를 계산 확정 로직과 분리해 설계", "중간"),
            ("QA", "청주 샘플처럼 기준자료가 불일치하는 케이스를 별도 테스트 시나리오로 고정", "중간"),
        ],
        [1700, 6160, 1500],
        body_size=9.1,
    )

    add_heading(doc, "10. 팀 공유 시 메시지", 1)
    add_callout(
        doc,
        "추천 설명",
        "공고문 기능을 버린 것이 아니라, 공고문을 만들기 전에 믿을 수 있는 원가검증 데이터부터 만드는 구조로 정리했습니다. Phase 1의 결과물인 validatedProject가 Phase 2 공고문 자동생성의 입력이 됩니다.",
        fill=WARNING_FILL,
    )

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_doc())
